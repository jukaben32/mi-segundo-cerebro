"""
convert_md_to_docx.py
Converts a Markdown (.md) file to a professionally formatted Word document (.docx).

Usage:
    python Scripts/convert_md_to_docx.py "path/to/file.md"
    python Scripts/convert_md_to_docx.py "path/to/file.md" "path/to/output.docx"

Requires: pip install python-docx
No external tools (Pandoc) needed.
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Palette ───────────────────────────────────────────────────────────────────
DARK_NAVY  = RGBColor(0x1B, 0x2A, 0x47)
ACCENT     = RGBColor(0x2E, 0x5F, 0x8A)
TEXT_DARK  = RGBColor(0x1C, 0x2B, 0x3A)
MID_GRAY   = RGBColor(0xB0, 0xBE, 0xC5)
LIGHT_GRAY = RGBColor(0xF4, 0xF6, 0xF8)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

BODY_FONT    = "Calibri"
HEADING_FONT = "Calibri"
BODY_SIZE    = Pt(11)
H1_SIZE      = Pt(20)
H2_SIZE      = Pt(15)
H3_SIZE      = Pt(12)


# ── Helpers ───────────────────────────────────────────────────────────────────

def strip_emojis(text: str) -> str:
    emoji_pattern = re.compile(
        "["
        u"\U0001F600-\U0001F64F"
        u"\U0001F300-\U0001F5FF"
        u"\U0001F680-\U0001F6FF"
        u"\U0001F1E0-\U0001F1FF"
        u"\U00002500-\U00002BEF"
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        u"\u2640-\u2642"
        u"\u2600-\u2B55"
        u"\u200d\u23cf\u23e9\u231a\ufe0f\u3030"
        "]+", flags=re.UNICODE
    )
    return emoji_pattern.sub("", text).strip()


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_paragraph(para, font_name=BODY_FONT, size=BODY_SIZE,
                     bold=False, color=TEXT_DARK, space_before=0, space_after=6):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    for run in para.runs:
        run.font.name  = font_name
        run.font.size  = size
        run.font.bold  = bold
        run.font.color.rgb = color


def add_inline_formatting(para, text: str, font_name=BODY_FONT,
                           size=BODY_SIZE, color=TEXT_DARK, bold_base=False):
    """Parse **bold** and `code` inline markers within a paragraph."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|`(.+?)`)")
    last = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if last < start:
            run = para.add_run(text[last:start])
            run.font.name = font_name
            run.font.size = size
            run.font.bold = bold_base
            run.font.color.rgb = color
        if match.group(1).startswith("**"):
            run = para.add_run(match.group(2))
            run.font.name = font_name
            run.font.size = size
            run.font.bold = True
            run.font.color.rgb = color
        else:
            run = para.add_run(match.group(3))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4C)
        last = end
    if last < len(text):
        run = para.add_run(text[last:])
        run.font.name = font_name
        run.font.size = size
        run.font.bold = bold_base
        run.font.color.rgb = color


def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)


def add_heading(doc, text: str, level: int):
    text = strip_emojis(text)
    para = doc.add_paragraph()
    if level == 1:
        run = para.add_run(text)
        run.font.name  = HEADING_FONT
        run.font.size  = H1_SIZE
        run.font.bold  = True
        run.font.color.rgb = DARK_NAVY
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after  = Pt(6)
        # Bottom border under H1
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1B2A47")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run = para.add_run(text)
        run.font.name  = HEADING_FONT
        run.font.size  = H2_SIZE
        run.font.bold  = True
        run.font.color.rgb = ACCENT
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run = para.add_run(text)
        run.font.name  = HEADING_FONT
        run.font.size  = H3_SIZE
        run.font.bold  = True
        run.font.color.rgb = TEXT_DARK
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after  = Pt(2)


def add_hr(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C5CDD6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)


def add_md_table(doc, header_row: list[str], rows: list[list[str]]):
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = strip_emojis(h.strip())
        set_cell_bg(hdr_cells[i], "2E5F8A")
        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name  = BODY_FONT
            run.font.size  = Pt(10)
            run.font.bold  = True
            run.font.color.rgb = WHITE
        table.rows[0].height = Pt(22)

    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F4F6F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row):
            cell_val = strip_emojis(cell_text.strip()) if c_idx < len(row) else ""
            row_cells[c_idx].text = cell_val
            set_cell_bg(row_cells[c_idx], bg)
            para = row_cells[c_idx].paragraphs[0]
            for run in para.runs:
                run.font.name  = BODY_FONT
                run.font.size  = Pt(10)
                run.font.color.rgb = TEXT_DARK
            table.rows[r_idx + 1].height = Pt(18)

    # Equal column widths
    col_width = Inches(6.2 / num_cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    doc.add_paragraph()  # spacing after table


# ── Main converter ────────────────────────────────────────────────────────────

def convert(md_path: str, output_path: str = None):
    md_path = Path(md_path)
    if not md_path.exists():
        print(f"Error: File not found — {md_path}")
        sys.exit(1)

    output_path = Path(output_path) if output_path else md_path.with_suffix(".docx")

    doc = Document()
    set_doc_margins(doc)

    # Default paragraph style baseline
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.font.color.rgb = TEXT_DARK

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code_block = False
    code_lines    = []

    while i < len(lines):
        line = lines[i]

        # ── Code block ───────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name  = "Courier New"
                run.font.size  = Pt(9.5)
                run.font.color.rgb = RGBColor(0x2E, 0x3A, 0x4E)
                para.paragraph_format.left_indent  = Cm(0.8)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after  = Pt(8)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()):
            add_hr(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text  = heading_match.group(2)
            add_heading(doc, text, level)
            i += 1
            continue

        # ── Markdown table ────────────────────────────────────────────────────
        if line.strip().startswith("|") and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if re.match(r"^\|[\s\-:|]+\|", next_line):
                header_row = [c.strip() for c in line.strip().strip("|").split("|")]
                i += 2  # skip separator
                rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    rows.append(row)
                    i += 1
                add_md_table(doc, header_row, rows)
                continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.strip().startswith(">"):
            text = re.sub(r"^>\s*(\[!.*?\])?\s*", "", line).strip()
            text = strip_emojis(text)
            if not text:
                i += 1
                continue
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Cm(0.8)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(text)
            run.font.name   = BODY_FONT
            run.font.size   = BODY_SIZE
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x62, 0x7A)
            i += 1
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text   = strip_emojis(bullet_match.group(2))
            para   = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Cm(0.4 + indent * 0.4)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline_formatting(para, text)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        num_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if num_match:
            indent = len(num_match.group(1))
            text   = strip_emojis(num_match.group(2))
            para   = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent  = Cm(0.4 + indent * 0.4)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline_formatting(para, text)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────────────
        text = strip_emojis(line)
        if not text:
            i += 1
            continue
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(6)
        add_inline_formatting(para, text)
        i += 1

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python convert_md_to_docx.py input.md [output.docx]")
        sys.exit(1)
    md_in    = sys.argv[1]
    docx_out = sys.argv[2] if len(sys.argv) > 2 else None
    convert(md_in, docx_out)
